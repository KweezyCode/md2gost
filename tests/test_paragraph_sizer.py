import unittest

import docx
from docx import Document

from md2gost.renderable.paragraph_sizer import Font, ParagraphSizer
from md2gost.renderable.listing import LISTING_OFFSET
from docx.shared import Pt, Mm, Cm

from . import _create_test_document, _EMUS_PER_PX

delta = 10/29


class TestFont(unittest.case.TestCase):
    def test_get_text_width(self):
        font = Font("Times New Roman", False, False, 14)
        self.assertAlmostEqual(37.5, font.get_text_width("hello") / _EMUS_PER_PX, delta=delta)

    def test_get_text_width_short(self):
        font = Font("Times New Roman", False, False, 14)
        self.assertAlmostEqual(15, font.get_text_width("in") / _EMUS_PER_PX, delta=delta)

    def test_get_text_width_long(self):
        font = Font("Times New Roman", False, False, 14)
        self.assertAlmostEqual(245, font.get_text_width("Электроэнцефалографический") / _EMUS_PER_PX, delta=delta)

    def test_get_text_width_bold(self):
        font = Font("Times New Roman", True, False, 14)
        self.assertAlmostEqual(39, font.get_text_width("hello") / _EMUS_PER_PX, delta=delta)

    def test_get_text_width_italic(self):
        font = Font("Times New Roman", False, True, 14)
        self.assertAlmostEqual(37.5, font.get_text_width("hello") / _EMUS_PER_PX, delta=delta)

    def test_get_text_width_bold_italic(self):
        font = Font("Times New Roman", True, True, 14)
        self.assertAlmostEqual(38.5, font.get_text_width("hello") / _EMUS_PER_PX, delta=delta)

    def test_get_line_height_times(self):
        font = Font("Times New Roman", False, False, 14)
        self.assertAlmostEqual(21.4, font.get_line_height() / _EMUS_PER_PX, delta=delta)

    def test_get_line_height_times_large(self):
        font = Font("Times New Roman", False, False, 50)
        self.assertAlmostEqual(77, font.get_line_height() / _EMUS_PER_PX, delta=delta)

    def test_get_line_height_calibri(self):
        font = Font("Calibri", False, False, 14)
        self.assertAlmostEqual(23, font.get_line_height() / _EMUS_PER_PX, delta=delta)

    def test_get_line_height_consolas(self):
        font = Font("Consolas", False, False, 20)
        self.assertAlmostEqual(31, font.get_line_height() / _EMUS_PER_PX, delta=delta)

    def test_get_line_height_courier(self):
        font = Font("Courier New", False, False, 12)
        self.assertAlmostEqual(18.3, font.get_line_height() / _EMUS_PER_PX, delta=delta)

    def test_is_mono_courier(self):
        font = Font("Courier New", False, False, 12)
        self.assertTrue(font.is_mono)

    def test_is_mono_times(self):
        font = Font("Times New Roman", False, False, 12)
        self.assertFalse(font.is_mono)


class TestParagraphSizer(unittest.TestCase):
    def setUp(self):
        self._document, self._max_height, self._max_width = _create_test_document()

    def test_count_lines(self):
        paragraph = self._document.add_paragraph()
        paragraph.add_run("Lorem ipsum dolor sit amet, consectetur adipiscing elit. Nam lacinia fringilla lectus, "
                          "nec euismod odio convallis sed. Nunc ac libero ultricies, condimentum neque et, "
                          "fermentum urna. Donec feugiat diam sed nulla rutrum, sit amet accumsan odio tempor. Sed "
                          "bibendum ante at orci faucibus, sed dignissim nisi finibus. Vestibulum luctus eget enim et "
                          "mattis. In porta convallis ipsum eget dignissim. Ut orci ante, bibendum ut lorem quis, "
                          "gravida molestie neque. Nulla vitae sapien sed risus gravida elementum non eu lorem. "
                          "Quisque ac turpis nisl.")

        ps = ParagraphSizer(paragraph, None, self._max_width)

        self.assertEqual(7, ps.count_lines(paragraph.runs, self._max_width, paragraph.style.font, Cm(1.25)))

    def test_count_lines2(self):
        paragraph = self._document.add_paragraph()
        paragraph.add_run("Ordered lists are useful when you want to present items in a specific order. This is additional text for illustration. Ordered lists are useful when you want to present items in a specific order. This is additional text for illustration. Ordered lists are useful when you want to present items in a specific order. This is additional text for illustration.")

        ps = ParagraphSizer(paragraph, None, self._max_width)

        self.assertEqual(4, ps.count_lines(paragraph.runs, self._max_width, paragraph.style.font, Cm(1.25)))

    def test_count_lines_short_last_line(self):
        paragraph = self._document.add_paragraph()
        paragraph.add_run("Nam porta urna vel turpis lobortis, nec congue sem suscipit. Mauris ac facilisis metus, "
                          "non fermentum turpis. Nunc erat ipsum, interdum sit amet odio ut, vulputate feugiat arcu. "
                          "Donec nec rhoncus metus, nec pharetra mauris. Proin magna arcu, porta vitae eleifend et, "
                          "scelerisque vel nunc. Ut ut neque sed libero mattis finibus quis eu tortor. Quisque "
                          "molestie tempus neque rutrum vestibulum. Proin dui odio, tincidunt nec lectus at, "
                          "mollis tristique diam. Nulla arcu ante, fringilla sit amet pretium venenatis, ultricies a "
                          "elit. Sed ac velit a dolor interdum sollicitudin. Donec tortor leo, finibus eu nisi id, "
                          "cursus luctus diam. Vestibulum ex nulla, fringilla pellentesque diam eu, rhoncus suscipit "
                          "ligula. Mauris vestibulum libero erat, vitae mollis orci ultricies at. Nulla maximus "
                          "elementum nulla at ultrices. Vestibulum pellentesque vulputate orci, quis finibus")

        ps = ParagraphSizer(paragraph, None, self._max_width)

        self.assertEqual(11, ps.count_lines(paragraph.runs, self._max_width, paragraph.style.font, Cm(1.25)))

    def test_count_lines_short_last_line2(self):
        paragraph = self._document.add_paragraph()
        paragraph.add_run("OpenAI is a leading artificial intelligence research organization, known for advancements in language models like GPT. Click the link to learn more. Hello world —")

        ps = ParagraphSizer(paragraph, None, self._max_width)

        self.assertEqual(3, ps.count_lines(paragraph.runs, self._max_width, paragraph.style.font, Cm(1.25)))

    def test_count_lines_long_last_line(self):
        paragraph = self._document.add_paragraph()
        paragraph.add_run("Donec finibus elementum lectus non ultricies. Pellentesque dictum tellus a neque rutrum "
                          "euismod. Fusce lobortis id est ut bibendum. Integer quis nunc convallis, maximus justo "
                          "fermentum, vestibulum metus. Nulla fringilla quam in purus laoreet, eu rhoncus risus "
                          "condimentum. Sed eget odio urna. Integer mi diam, aliquam id rhoncus vitae, lacinia quis "
                          "augue. Nulla ultrices velit vel urna accumsan, et feugiat nunc fringilla. Praesent feugiat "
                          "neque ac tellus rutrum congue. Sed congue libero congue, blandit felis ac, "
                          "lobortis libero. Phasellus eleifend ex vulputate odio mollis dictum.")

        ps = ParagraphSizer(paragraph, None, self._max_width)

        self.assertEqual(7, ps.count_lines(paragraph.runs, self._max_width, paragraph.style.font, Cm(1.25)))

    def test_count_lines_long_last_line2(self):
        paragraph = self._document.add_paragraph()
        paragraph.add_run('Markdown supports rendering mathematical formulas using LaTeX syntax. This allows you to include complex equations and mathematical notation in your documents.')

        ps = ParagraphSizer(paragraph, None, self._max_width)

        self.assertEqual(2, ps.count_lines(paragraph.runs, self._max_width, paragraph.style.font, Cm(1.25)))

    def test_count_lines_long_last_line3(self):
        paragraph = self._document.add_paragraph()
        paragraph.add_run("Markdown's LaTeX syntax allows you to easily write fractions and exponents. For instance, you can represent the derivative of a function  InlineFormula is not supported  with respect to  InlineFormula is not supported  using the following notation:")

        ps = ParagraphSizer(paragraph, None, self._max_width)

        self.assertEqual(3, ps.count_lines(paragraph.runs, self._max_width, paragraph.style.font, Cm(1.25)))


    def test_count_lines_long_word(self):
        paragraph = self._document.add_paragraph()
        paragraph.add_run("verylongwordverylongwordverylongwordverylongwordverylongwordverylongwordverylongwordverylongwordverylongwordverylongwordverylongwordverylongwordverylongword")

        ps = ParagraphSizer(paragraph, None, self._max_width)

        self.assertEqual(3, ps.count_lines(paragraph.runs, self._max_width, paragraph.style.font, Cm(1.25)))

    def test_count_lines_long_word_2(self):
        paragraph = self._document.add_paragraph()
        paragraph.add_run("someword verylongwordverylongwordverylongwordverylongwordverylongwordverylongwordverylongwordverylongwordverylongwordverylongwordverylongwordverylongwordverylongword")

        ps = ParagraphSizer(paragraph, None, self._max_width)

        self.assertEqual(4, ps.count_lines(paragraph.runs, self._max_width, paragraph.style.font, Cm(1.25)))

    def test_count_lines_courier(self):
        paragraph = self._document.add_paragraph(style="Code")
        paragraph.add_run("""            table._cells[0]._element.append(paragraph_rendered_info.docx_element._element)""")

        ps = ParagraphSizer(paragraph, None, self._max_width-LISTING_OFFSET)

        self.assertEqual(3, ps.count_lines(paragraph.runs, self._max_width-LISTING_OFFSET, paragraph.style.font, 0))

    def test_count_lines_courier2(self):
        paragraph = self._document.add_paragraph(style="Code")
        paragraph.add_run(
            """                continuation_paragraph = Paragraph(self.parent)""")

        ps = ParagraphSizer(paragraph, None, self._max_width - LISTING_OFFSET)

        self.assertEqual(1, ps.count_lines(paragraph.runs, self._max_width - LISTING_OFFSET, paragraph.style.font, 0))

    def test_count_lines_courier3(self):
        paragraph = self._document.add_paragraph(style="Code")
        paragraph.add_run(
            """ooooooooooooooooooooooooooooooooooooooooooooooooooooooooooooooooooooo""")

        ps = ParagraphSizer(paragraph, None, self._max_width - Pt(14))

        self.assertEqual(2, ps.count_lines(paragraph.runs, self._max_width - Pt(14), paragraph.style.font, 0))

    def test_count_lines_courier4(self):
        paragraph = self._document.add_paragraph(style="Code")
        paragraph.add_run(
            """    def add_link(self, text: str, url: str, is_bold: bool = None, is_italic: bool = None):""")

        ps = ParagraphSizer(paragraph, None, self._max_width - LISTING_OFFSET)

        self.assertEqual(2, ps.count_lines(paragraph.runs, self._max_width - LISTING_OFFSET, paragraph.style.font, 0))

    def test_count_lines_courier5(self):
        paragraph = self._document.add_paragraph(style="Code")
        paragraph.add_run(
            """                self._docx_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY""")

        ps = ParagraphSizer(paragraph, None, self._max_width - LISTING_OFFSET)

        self.assertEqual(2, ps.count_lines(paragraph.runs, self._max_width - LISTING_OFFSET, paragraph.style.font, 0))

    def test_count_lines_courier6(self):
        paragraph = self._document.add_paragraph(style="Code")
        paragraph.add_run(
            """            #         and "Heading" in previous_rendered.docx_element.style.name\\""")

        ps = ParagraphSizer(paragraph, None, self._max_width - LISTING_OFFSET)

        self.assertEqual(2, ps.count_lines(paragraph.runs, self._max_width - LISTING_OFFSET, paragraph.style.font, 0))

    def test_count_lines_courier7(self):
        paragraph = self._document.add_paragraph(style="Code")
        paragraph.add_run(
            """        run = DocxRun(create_element("w:r"), self._docx_paragraph)""")

        ps = ParagraphSizer(paragraph, None, self._max_width - LISTING_OFFSET)

        self.assertEqual(1, ps.count_lines(paragraph.runs, self._max_width - LISTING_OFFSET, paragraph.style.font, 0, True))

    def test_count_lines_courier8(self):
        paragraph = self._document.add_paragraph(style="Code")
        paragraph.add_run(
            """            #         and ((min(2, height_data.lines) - 1) * height_data.line_spacing + 1) * height_data.line_height\\""")

        ps = ParagraphSizer(paragraph, None, self._max_width - LISTING_OFFSET)

        self.assertEqual(2, ps.count_lines(paragraph.runs, self._max_width - LISTING_OFFSET, paragraph.style.font, 0, True))

    def test_count_lines_courier9(self):
        paragraph = self._document.add_paragraph(style="Code")
        paragraph.add_run(
            """        return self._docx_paragraph.paragraph_format.first_line_indent""")

        ps = ParagraphSizer(paragraph, None, self._max_width - LISTING_OFFSET)

        self.assertEqual(2, ps.count_lines(paragraph.runs, self._max_width - LISTING_OFFSET, paragraph.style.font, 0, True))

    def test_count_lines_courier10(self):
        paragraph = self._document.add_paragraph(style="Code")
        paragraph.add_run(
            """        self._docx_paragraph.paragraph_format.first_line_indent = value""")

        ps = ParagraphSizer(paragraph, None, self._max_width - LISTING_OFFSET)

        self.assertEqual(2, ps.count_lines(paragraph.runs, self._max_width - LISTING_OFFSET, paragraph.style.font, 0, True))

    def test_count_lines_courier11(self):
        paragraph = self._document.add_paragraph(style="Code")
        paragraph.add_run(
            """                if previous_rendered and isinstance(previous_rendered.docx_element, DocxParagraph) else None,""")

        ps = ParagraphSizer(paragraph, None, self._max_width - LISTING_OFFSET)

        self.assertEqual(2, ps.count_lines(paragraph.runs, self._max_width - LISTING_OFFSET, paragraph.style.font, 0, True))

    def test_count_lines_courier12(self):
        paragraph = self._document.add_paragraph(style="Code")
        paragraph.add_run(
            '        System.out.printf("Square: %f%n", circle.calculateSquare());')

        ps = ParagraphSizer(paragraph, None, self._max_width - LISTING_OFFSET)

        self.assertEqual(1, ps.count_lines(paragraph.runs, self._max_width - LISTING_OFFSET, paragraph.style.font, 0, True))

    def test_count_lines_courier13(self):
        paragraph = self._document.add_paragraph(style="Code")
        paragraph.add_run(
            '        System.out.printf("circle radius: %f%n", circle.getRadius());')

        ps = ParagraphSizer(paragraph, None, self._max_width - LISTING_OFFSET)

        self.assertEqual(2, ps.count_lines(paragraph.runs, self._max_width - LISTING_OFFSET, paragraph.style.font, 0, True))


    def test_count_lines_courier_multiple_runs(self):
        paragraph = self._document.add_paragraph(style="Code")
        for run_text in ['', '        ', 'run', ' ', '=', ' ', 'DocxRun', '(', 'create_element', '(', '"', 'w:r', '"', ')', ',', ' ', 'self', '.', '_docx_paragraph', ')']:
            paragraph.add_run(run_text)

        ps = ParagraphSizer(paragraph, None, self._max_width - LISTING_OFFSET)

        self.assertEqual(1, ps.count_lines(paragraph.runs, self._max_width - LISTING_OFFSET, paragraph.style.font, 0, True))

    def test_count_lines_courier_multiple_runs2(self):
        paragraph = self._document.add_paragraph(style="Code")
        for run_text in ['', '            ', '-', '>', ' ', 'Generator', '[', 'RenderedInfo', ' ', '|', ' ', 'Renderable', ',', ' ', 'None', ',', ' ', 'None', ']', ':']:
            paragraph.add_run(run_text)

        ps = ParagraphSizer(paragraph, None, self._max_width - LISTING_OFFSET)

        self.assertEqual(1, ps.count_lines(paragraph.runs, self._max_width - LISTING_OFFSET, paragraph.style.font, 0, True))    

    # def test_count_lines_courier_multiple_runs3(self):
    #     paragraph = self._document.add_paragraph(style="Code")
    #     for run_text in ['', '            ', '-', '>', ' ', 'Generator', '[', 'RenderedInfo', ' ', '|', ' ', 'Renderable', ',', ' ', 'None', ',', ' ', 'None', ']', ':']:
    #         paragraph.add_run(run_text)
    # 
    #     ps = ParagraphSizer(paragraph, None, self._max_width - LISTING_OFFSET)
    # 
    #     self.assertEqual(3, ps.count_lines(paragraph.runs, self._max_width - LISTING_OFFSET, paragraph.style.font, 0, True))
